#!/usr/bin/env python3
"""
Generate Register of Directors (ROD) and Register of Members (ROM) as .docx files.
Usage: python generate_registers.py <company_json_file> <type: rod|rom> <output_path>
Reads company JSON from stdin if no file path provided.
"""
import json
import sys
import os
from datetime import datetime

sys.path.insert(0, '/tmp/pylibs')
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def parse_date(date_val):
    """Parse date string to YYYY-MM-DD or return as-is."""
    if not date_val:
        return 'Present'
    try:
        if isinstance(date_val, str):
            return date_val[:10] if 'T' in date_val else date_val[:10] if len(date_val) >= 10 else date_val
        return str(date_val)[:10]
    except:
        return str(date_val)


def set_cell_border(cell, **kwargs):
    """Set cell border."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = tcPr.makeelement(qn('w:tcBorders'), {})
        tcPr.append(tcBorders)
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        if edge in kwargs:
            element = tcBorders.find(qn(f'w:{edge}'))
            if element is None:
                element = tcBorders.makeelement(qn(f'w:{edge}'), {})
                tcBorders.append(element)
            for attr, val in kwargs[edge].items():
                element.set(qn(f'w:{attr}'), str(val))


def add_bordered_table(doc, headers, rows, col_widths=None):
    """Add a bordered table to the document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set column widths
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    # Header row
    hdr = table.rows[0]
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        # Grey background
        shading = cell._tc.get_or_add_tcPr()
        shd = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): 'D9D9D9',
            qn('w:val'): 'clear',
        })
        shading.append(shd)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(text) if text else '')
            run.font.size = Pt(8.5)
            run.font.name = 'Times New Roman'

    return table


def generate_rod(company_data, output_path):
    """Generate Register of Directors."""
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(company_data.get('name', 'Company Name'))
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    if company_data.get('nameChinese'):
        cn = doc.add_paragraph()
        cn.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cn.add_run(company_data['nameChinese'])
        run.font.size = Pt(11)
        run.font.name = 'Microsoft YaHei'

    # Register Title
    reg_title = doc.add_paragraph()
    reg_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    reg_title.paragraph_format.space_before = Pt(12)
    reg_title.paragraph_format.space_after = Pt(12)
    run = reg_title.add_run('REGISTER OF DIRECTORS')
    run.bold = True
    run.font.size = Pt(13)
    run.font.name = 'Times New Roman'

    # Meta line
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(16)
    run = meta.add_run(
        f"Company No.: {company_data.get('registrationNumber', 'N/A')}    |    "
        f"Jurisdiction: {company_data.get('jurisdiction', 'N/A')}    |    "
        f"Date: {datetime.now().strftime('%Y-%m-%d')}"
    )
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'

    # Table
    links = company_data.get('links', [])
    directors = [l for l in links if 'director' in l.get('roles', []) or 'alternate_director' in l.get('roles', [])]

    headers = ['Date Appointed', 'Full Name', 'NRIC / Passport', 'Nationality', 'Address', 'Role', 'Date Ceased']
    col_widths = [2.2, 3.5, 2.2, 1.8, 3.5, 2.0, 2.0]

    rows = []
    for d in directors:
        p = d.get('link', {})
        role_str = ', '.join(d.get('roles', []))
        rows.append([
            parse_date(d.get('appointedDate')),
            p.get('name', ''),
            p.get('nric', ''),
            p.get('nationality', ''),
            p.get('address', {}).get('country', '') if isinstance(p.get('address'), dict) else '',
            role_str,
            parse_date(d.get('ceasedDate')),
        ])

    if not rows:
        rows.append(['No directors registered', '', '', '', '', '', ''])

    add_bordered_table(doc, headers, rows, col_widths)

    # Footer
    footer = doc.add_paragraph()
    footer.paragraph_format.space_before = Pt(30)
    run = footer.add_run(f'Generated by Claw Company Secretary System on {datetime.now().strftime("%Y-%m-%d")}')
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(128, 128, 128)
    run.font.name = 'Times New Roman'

    doc.save(output_path)
    return output_path


def generate_rom(company_data, output_path):
    """Generate Register of Members."""
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(company_data.get('name', 'Company Name'))
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    if company_data.get('nameChinese'):
        cn = doc.add_paragraph()
        cn.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cn.add_run(company_data['nameChinese'])
        run.font.size = Pt(11)
        run.font.name = 'Microsoft YaHei'

    # Register Title
    reg_title = doc.add_paragraph()
    reg_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    reg_title.paragraph_format.space_before = Pt(12)
    reg_title.paragraph_format.space_after = Pt(12)
    run = reg_title.add_run('REGISTER OF MEMBERS')
    run.bold = True
    run.font.size = Pt(13)
    run.font.name = 'Times New Roman'

    # Meta line
    share_cap = company_data.get('shareCapital', {})
    issued = share_cap.get('issued', 0)
    currency = share_cap.get('currency', 'HKD')

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(16)
    run = meta.add_run(
        f"Company No.: {company_data.get('registrationNumber', 'N/A')}    |    "
        f"Jurisdiction: {company_data.get('jurisdiction', 'N/A')}    |    "
        f"Issued Shares: {issued:,} {currency}    |    "
        f"Date: {datetime.now().strftime('%Y-%m-%d')}"
    )
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'

    # Table
    links = company_data.get('links', [])
    shareholders = [l for l in links if 'shareholder' in l.get('roles', [])]

    headers = ['Date Entered', 'Member Name', 'Address / Jurisdiction', 'No. of Shares', 'Share Type', 'Shareholding %', 'Date Ceased']
    col_widths = [2.2, 3.5, 3.0, 1.8, 1.6, 1.8, 2.0]

    rows = []
    total_shares = issued or 1
    for s in shareholders:
        p = s.get('link', {})
        shares = s.get('shares', 0) or 0
        pct = f"{(shares / total_shares * 100):.2f}%" if total_shares > 0 else '-'
        rows.append([
            parse_date(s.get('appointedDate')),
            p.get('name', ''),
            p.get('address', {}).get('country', '') if isinstance(p.get('address'), dict) else p.get('registrationNumber', ''),
            f"{shares:,}",
            s.get('shareType', 'Ordinary'),
            pct,
            parse_date(s.get('ceasedDate')),
        ])

    if not rows:
        rows.append(['No members registered', '', '', '', '', '', ''])

    add_bordered_table(doc, headers, rows, col_widths)

    # Footer
    footer = doc.add_paragraph()
    footer.paragraph_format.space_before = Pt(30)
    run = footer.add_run(f'Generated by Claw Company Secretary System on {datetime.now().strftime("%Y-%m-%d")}')
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(128, 128, 128)
    run.font.name = 'Times New Roman'

    doc.save(output_path)
    return output_path


def generate_register_from_file(json_path, reg_type, output_path):
    """Load company JSON from file and generate register."""
    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    if reg_type == 'rod':
        return generate_rod(data, output_path)
    elif reg_type == 'rom':
        return generate_rom(data, output_path)
    else:
        raise ValueError(f"Unknown register type: {reg_type}")


def generate_register_from_stdin(reg_type, output_path):
    """Read company JSON from stdin and generate register."""
    data = json.load(sys.stdin)
    if reg_type == 'rod':
        return generate_rod(data, output_path)
    elif reg_type == 'rom':
        return generate_rom(data, output_path)
    else:
        raise ValueError(f"Unknown register type: {reg_type}")


if __name__ == '__main__':
    n = len(sys.argv)
    if n == 2:
        # stdin mode: reg_type only, default output
        generate_register_from_stdin(sys.argv[1], f"output_{sys.argv[1]}.docx")
    elif n == 3:
        # stdin mode: reg_type + output_path
        generate_register_from_stdin(sys.argv[1], sys.argv[2])
    elif n == 4:
        # file mode: json_file + reg_type + output_path
        generate_register_from_file(sys.argv[1], sys.argv[2], sys.argv[3])
    else:
        print("Usage: echo '<json>' | python generate_registers.py <type: rod|rom> [output_path]", file=sys.stderr)
        print("       python generate_registers.py <json_file> <type: rod|rom> <output_path>", file=sys.stderr)
        sys.exit(1)
